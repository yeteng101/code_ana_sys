from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path('/Users/andye/Library/Containers/com.tencent.xinWeChat/Data/Documents/app_data/radium/users/7d84a21d2e62569ab152d8e2a06c48bf/applet/local/wx6bccbeb9af7d80bf/temp/孙业腾_简历.docx')
TARGET = Path('/Users/andye/Documents/ChatGPT/8.18huawei/孙业腾_简历-已加入代码逆向Agent.docx')


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = 0


def clone_paragraph_with_text(template, text):
    paragraph_xml = deepcopy(template._p)
    paragraph = template._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    paragraph._p.addnext(paragraph_xml)
    paragraph = template._parent.paragraphs[template._parent.paragraphs.index(template) + 1]
    # Preserve paragraph properties and the first run's formatting.
    for child in list(paragraph_xml):
        if child.tag != qn('w:pPr'):
            paragraph_xml.remove(child)
    run = OxmlElement('w:r')
    template_run = template.runs[0] if template.runs else None
    if template_run is not None and template_run._r.rPr is not None:
        run.append(deepcopy(template_run._r.rPr))
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    paragraph_xml.append(run)
    return paragraph


def insert_before(parent, new_element, reference_element):
    reference_element.addprevious(new_element)


doc = Document(str(SOURCE))
body = doc.element.body

# The first project table is the table immediately after the 项目经历 heading.
project_table = doc.tables[2]
project_paragraph = next(p for p in doc.paragraphs if p.text.startswith('•  项目背景：'))

new_table_xml = deepcopy(project_table._tbl)
project_table._tbl.addprevious(new_table_xml)
new_table = doc.tables[2]
set_cell_text(new_table.cell(0, 0), '代码逆向 Agent 设计与 Demo')
set_cell_text(new_table.cell(0, 1), '实习项目')
set_cell_text(new_table.cell(0, 2), '华为数据通信产品线·软转发与 NaaS 部门')

bullets = [
    '•  项目背景：参与华为数据通信产品线软转发与 NaaS 部门实习项目，面向 libuv、Redis 等开源 C/C++ 代码，探索自动逆向分析模块架构、关键函数调用链和异步回调链。',
    '•  架构与规范：设计由源码索引、编译宏分析、调用图、函数指针、异步事件和验证 Agent 组成的 Subagent 架构，定义任务编排、证据链和置信度规范。',
    '•  源码分析：梳理 libuv 事件循环与 handle 生命周期、Redis 事件驱动请求链，区分直接调用、回调边、函数指针候选和跨平台实现。',
    '•  工程产出：设计 AnalysisRequest / AgentResult JSON Schema，输出 Mermaid 调用关系图、Demo 验证报告和汇报 PPT，强调固定 commit、编译配置和源码位置可复核。',
]

# Insert the new project bullets between the new table and the existing HTTP project table.
reference_table = project_table._tbl
for text in bullets:
    paragraph_xml = deepcopy(project_paragraph._p)
    for child in list(paragraph_xml):
        if child.tag != qn('w:pPr'):
            paragraph_xml.remove(child)
    run = OxmlElement('w:r')
    if project_paragraph.runs and project_paragraph.runs[0]._r.rPr is not None:
        run.append(deepcopy(project_paragraph.runs[0]._r.rPr))
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    paragraph_xml.append(run)
    reference_table.addprevious(paragraph_xml)

doc.save(str(TARGET))
print(TARGET)
