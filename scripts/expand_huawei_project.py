from pathlib import Path

from docx import Document


path = Path('/Users/andye/Documents/ChatGPT/8.18huawei/孙业腾_简历-已加入代码逆向Agent.docx')
doc = Document(str(path))

updates = {
    '•  项目背景：参与华为数据通信产品线软转发与 NaaS 部门实习项目，面向 libuv、Redis 等开源 C/C++ 代码，探索自动逆向分析模块架构、关键函数调用链和异步回调链。':
    '•  项目背景：参与华为数据通信产品线软转发与 NaaS 部门实习项目，围绕 libuv、Redis 等开源 C/C++ 基础组件开展源码逆向分析，梳理事件循环、网络 I/O、客户端请求处理等模块边界，重点还原关键函数调用链、异步回调链和句柄生命周期，为架构理解与问题定位提供结构化分析基础。',
    '•  架构与规范：设计由源码索引、编译宏分析、调用图、函数指针、异步事件和验证 Agent 组成的 Subagent 架构，定义任务编排、证据链和置信度规范。':
    '•  架构与规范：设计“源码索引—编译宏分析—调用图—函数指针—异步事件—验证”多 Agent 协作架构，拆分仓库解析、跨文件调用追踪、回调注册/触发关联和平台实现识别等任务；定义统一 JSON 接口、Evidence Graph、源码位置引用、配置条件与置信度字段，支持结果合并、冲突校验和自然语言报告生成。',
    '•  研究问题：《ECG-aBcDe: Encoding ECG into a Universal Language for any Large Language Model.》探索将连续 ECG 信号编码为与具体 LLM 解耦的通用语言表示，使不同大语言模型能够直接接入并理解 ECG 信息。':
    '•  研究问题：论文《ECG-aBcDe: Overcoming model dependence, encoding ECG into a universal language for any large language model》面向 ECG 与 LLM 融合中的模型依赖、时间尺度建模和结果可解释性问题，探索可跨模型复用的 ECG 表示。',
    '•  问题分析：针对现有“两阶段/端到端”方案中 ECG 编码器与目标 LLM 强绑定、跨模型迁移成本高的问题，参与梳理通用编码框架及其与 LLM 的接口关系，明确“一次编码、多模型复用”的设计目标。':
    '•  方法设计：通过检测 ECG 关键点并量化关键点间间隔，将连续信号编码为交替包含波形关键点与时间间隔的离散“ECG language”；构造 ECG language 与自然语言混合数据集，在不修改 LLM 架构的情况下直接微调预训练 LLM。',
    '•  主要贡献：论文提出与 LLM 无关的 ECG 编码方法，不依赖特定 LLM 重新训练 ECG 编码器；基于统一表示支持跨模型复用，降低模型迁移与部署成本。':
    '•  实验结果：在 PTB-XL / MIMIC-IV 的同分布和零样本跨数据集任务上验证模型适配性；BLEU-4 最高 43.37，跨数据集最高 35.82；微调显存降低 36.4%，PTB-XL 诊断任务平均准确率达到 72.18%。',
    '•  个人贡献：作为学生第二作者，参与论文阅读、问题建模、创新点提炼与实验结果分析，协助将研究动机、方法价值和实验结论转化为结构化表达。':
    '•  个人贡献：作为第三作者，参与数据整理、软件实现、实验考察和结果可视化，协助完成实验材料整理与论文结果表达。',
    '学生第二作者':
    '学生第三作者'
}

def iter_paragraphs(container):
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


changed = 0
for paragraph in iter_paragraphs(doc):
    old_text = paragraph.text
    if old_text in updates and paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ''
        paragraph.runs[0].text = updates[old_text]
        changed += 1

missing = [text for text in updates.values()
           if not any(p.text == text for p in iter_paragraphs(doc))]
if missing:
    raise RuntimeError(f'updated paragraphs missing from document: {missing}')

doc.save(str(path))
print(path)
